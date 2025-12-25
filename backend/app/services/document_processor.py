"""
Document processing service for extracting text from DOCX and PDF files.
Handles file validation, text extraction, and chunking.
"""

import os
import tempfile
from pathlib import Path
from typing import List, Tuple
import docx
import PyPDF2
from fastapi import UploadFile, HTTPException


class DocumentProcessor:
    """Service for processing uploaded documents - UOIONHHC"""
    
    @staticmethod
    def validate_file(file: UploadFile, max_size: int = 10 * 1024 * 1024) -> None:
        """
        Validate uploaded file.
        
        Args:
            file: Uploaded file
            max_size: Maximum file size in bytes
            
        Raises:
            HTTPException if validation fails
        """
        # Check file extension
        file_ext = Path(file.filename).suffix.lower()
        if file_ext not in ['.pdf', '.docx', '.doc']:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid file type: {file_ext}. Only .pdf and .docx files are allowed."
            )
        
        # Check file size
        file.file.seek(0, 2)  # Seek to end
        file_size = file.file.tell()
        file.file.seek(0)  # Reset to beginning
        
        if file_size > max_size:
            raise HTTPException(
                status_code=400,
                detail=f"File too large: {file_size} bytes. Maximum size is {max_size} bytes."
            )
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        try:
            doc = docx.Document(file_path)
            text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            return '\n\n'.join(text)
        
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error extracting text from DOCX: {str(e)}"
            )
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            text = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text.append(page_text)
            
            return '\n\n'.join(text)
        
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Error extracting text from PDF: {str(e)}"
            )
    
    @staticmethod
    async def extract_text(file: UploadFile) -> Tuple[str, str]:
        """
        Extract text from uploaded file.
        
        Args:
            file: Uploaded file
            
        Returns:
            Tuple of (extracted_text, file_path)
        """
        # Save to temp file
        file_ext = Path(file.filename).suffix.lower()
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_path = temp_file.name
        
        try:
            # Extract based on file type
            if file_ext == '.pdf':
                text = DocumentProcessor.extract_text_from_pdf(temp_path)
            elif file_ext in ['.docx', '.doc']:
                text = DocumentProcessor.extract_text_from_docx(temp_path)
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {file_ext}"
                )
            
            return text, temp_path
        
        except Exception as e:
            # Clean up temp file
            if os.path.exists(temp_path):
                os.unlink(temp_path)
            raise e
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 4000, overlap: int = 200) -> List[str]:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Input text
            chunk_size: Size of each chunk in characters
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundaries
            if end < len(text):
                # Look for sentence end in last 100 chars
                search_start = max(end - 100, start)
                sentence_end = text.rfind('.', search_start, end)
                
                if sentence_end > start:
                    end = sentence_end + 1
            
            chunks.append(text[start:end])
            start = end - overlap
        
        return chunks
    
    @staticmethod
    def create_markdown_template(
        text: str,
        variables: List[dict],
        metadata: dict
    ) -> str:
        """
        Create Markdown template with YAML front-matter.
        
        Args:
            text: Template body text
            variables: Variable definitions
            metadata: Template metadata
            
        Returns:
            Markdown string with front-matter
        """
        import yaml
        
        # Build front-matter
        front_matter = {
            'template_id': metadata.get('template_id', ''),
            'title': metadata.get('title', ''),
            'file_description': metadata.get('file_description', ''),
            'jurisdiction': metadata.get('jurisdiction', ''),
            'doc_type': metadata.get('doc_type', ''),
            'variables': variables,
            'similarity_tags': metadata.get('similarity_tags', [])
        }
        
        # Convert to YAML
        yaml_str = yaml.dump(front_matter, default_flow_style=False, allow_unicode=True)
        
        # Combine with body
        template = f"---\n{yaml_str}---\n\n{text}"
        
        return template


# Global instance
document_processor = DocumentProcessor()
